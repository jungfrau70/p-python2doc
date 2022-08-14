
import docx
import re  
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()
sections = document.sections

for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
section = document.sections[0]

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '1')
    
def addSummaryTable(df, regex, month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=3, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '내용 상세'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['작업내용'][ind]  


def addTable3(df, regex, month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=3, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '제목'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['제목'][ind] 
        
def addTablek8s(df, regex, month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass

    unique_df = df['클러스터명'].value_counts()
    unique_df = pd.DataFrame({'클러스터명':unique_df.index, '노드수':unique_df.values})
    df = df.drop_duplicates(['클러스터명'])
    table = document.add_table(rows=1, cols=4, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '클러스터명'
    hdr_cells[2].text = '노드수'
    hdr_cells[3].text = '프로메테우스_설치_유무'
    
    
    i = 0
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = df['클러스터명'][ind]
        row_cells[2].text = str(unique_df['노드수'][i])
        row_cells[3].text = df['프로메테우스_설치_유무'][ind] 
        i += 1
        
def addTable4(df, regex, month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=4, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '제목'
    hdr_cells[3].text = '작업내용'    

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['제목'][ind]
        row_cells[3].text = df['작업내용'][ind]  
        

def addPivot2Table(df):
    
    table = document.add_table(rows=1, cols=9, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '유형'
    hdr_cells[1].text = 'KR'
    hdr_cells[2].text = 'NA'
    hdr_cells[3].text = 'EU'    
    hdr_cells[4].text = 'CN'
    hdr_cells[5].text = 'SG'
    hdr_cells[6].text = 'RU'
    hdr_cells[7].text = '합계' 
    hdr_cells[8].text = '비중' 
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['유형'][ind]
        row_cells[1].text = '%0.1f' % df['KR'][ind]
        row_cells[2].text = '%0.1f' % df['NA'][ind]
        row_cells[3].text = '%0.1f' % df['EU'][ind]  
        row_cells[4].text = '%0.1f' % df['CN'][ind]
        row_cells[5].text = '%0.1f' % df['SG'][ind]
        row_cells[6].text = '%0.1f' % df['RU'][ind]
        row_cells[7].text = '%0.1f' % df['합계'][ind]  
        row_cells[8].text = '%0.1f' % df['비중'][ind]  
        
def addFailedTable(df, regex, month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=5, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '날짜'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '장애이벤트'    
    hdr_cells[3].text = '조치내용'
    hdr_cells[4].text = '장애전파소요시간(분)'
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['날짜'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['장애이벤트'][ind]
        row_cells[3].text = df['조치내용'][ind]  
        row_cells[4].text = '%0.1f' % df['장애전파소요시간'][ind]

def addTable_DB_license(df, regex, month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=7, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '년도'
    hdr_cells[1].text = '월'
    hdr_cells[2].text = '리전'
    hdr_cells[3].text = '라이선스 종류'    
    hdr_cells[4].text = 'Hostname'    
    hdr_cells[5].text = '라이선스'    
    hdr_cells[6].text = '만료일'    

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['년도'][ind]
        row_cells[1].text = df['월'][ind]
        row_cells[2].text = df['리전'][ind]
        row_cells[3].text = df['라이선스 종류'][ind]
        row_cells[4].text = df['Hostname'][ind]          
        row_cells[5].text = df['라이선스'][ind]          
        row_cells[6].text = df['비고'][ind]
        
def addTable(df):
       
    table = document.add_table(df.shape[0]+1, df.shape[1], style="Table Grid")

    # add the header rows
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]

    # add the rest of the dataframe    
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])
